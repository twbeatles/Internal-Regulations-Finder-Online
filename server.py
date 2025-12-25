# -*- coding: utf-8 -*-
"""
ì‚¬ë‚´ ê·œì • ê²€ìƒ‰ê¸° - ì›¹ ì„œë²„ ë²„ì „ v1.0
Flask ê¸°ë°˜ ì›¹ ì„œë²„ë¡œ ë‹¤ì¤‘ ì‚¬ìš©ì ë™ì‹œ ì ‘ì† ì§€ì›
"""

from __future__ import annotations
import sys
import os
import json
import threading
import tempfile
import hashlib
import shutil
import logging
import subprocess
import platform
import re
import gc
import math
import time
from typing import List, Dict, Tuple, Optional, Any
from datetime import datetime
from dataclasses import dataclass, field
from enum import Enum, auto
from collections import Counter
from functools import lru_cache
from concurrent.futures import ThreadPoolExecutor
import signal
import atexit

from flask import Flask, request, jsonify, render_template, send_from_directory, session
from flask_cors import CORS
from werkzeug.utils import secure_filename

# ============================================================================
# ìƒìˆ˜ ë° ì„¤ì •
# ============================================================================
class AppConfig:
    APP_NAME = "ì‚¬ë‚´ ê·œì • ê²€ìƒ‰ê¸°"
    APP_VERSION = "1.7 (ì›¹ ì„œë²„)"  # v1.7 ë””ë²„ê¹… ë° ìµœì í™”
    
    # ì„œë²„ ì„¤ì •
    SERVER_HOST = "0.0.0.0"
    SERVER_PORT = 8080
    MAX_CONTENT_LENGTH = 50 * 1024 * 1024  # 50MB
    
    # ì˜¤í”„ë¼ì¸ ëª¨ë“œ ì„¤ì • (íì‡„ë§ ì§€ì›)
    OFFLINE_MODE = False  # Trueë©´ ì¸í„°ë„· ì—°ê²° ì—†ì´ ë¡œì»¬ ëª¨ë¸ë§Œ ì‚¬ìš©
    LOCAL_MODEL_PATH = ""  # ì‚¬ì „ ë‹¤ìš´ë¡œë“œëœ ëª¨ë¸ í´ë” ê²½ë¡œ (ë¹ˆ ë¬¸ìì—´ì´ë©´ ê¸°ë³¸ ê²½ë¡œ ì‚¬ìš©)
    
    # AI ëª¨ë¸ ì„¤ì •
    AVAILABLE_MODELS: Dict[str, str] = {
        "SNU SBERT (ê³ ì„±ëŠ¥)": "snunlp/KR-SBERT-V40K-klueNLI-augSTS",
        "BM-K Simal (ê· í˜•)": "BM-K/ko-simal-roberta-base",
        "JHGan SBERT (ë¹ ë¦„)": "jhgan/ko-sbert-nli"
    }
    DEFAULT_MODEL = "SNU SBERT (ê³ ì„±ëŠ¥)"
    
    # íŒŒì¼ ì„¤ì •
    UPLOAD_FOLDER = "uploads"
    SUPPORTED_EXTENSIONS = {'.txt', '.docx', '.pdf'}
    
    # ê²€ìƒ‰ ì„¤ì •
    MAX_SEARCH_RESULTS = 10
    DEFAULT_SEARCH_RESULTS = 5
    
    # ì²­í‚¹ ì„¤ì •
    CHUNK_SIZE = 800
    CHUNK_OVERLAP = 80
    VECTOR_WEIGHT = 0.7
    BM25_WEIGHT = 0.3
    
    # ë™ì‹œì„± ì„¤ì •
    MAX_WORKERS = 8
    REQUEST_TIMEOUT = 30
    SEARCH_CACHE_SIZE = 200
    MAX_CONCURRENT_SEARCHES = 10
    RATE_LIMIT_PER_MINUTE = 60


class FileStatus(Enum):
    PENDING = "ëŒ€ê¸°"
    PROCESSING = "ì²˜ë¦¬ì¤‘"
    SUCCESS = "ì™„ë£Œ"
    FAILED = "ì‹¤íŒ¨"
    CACHED = "ìºì‹œ"


@dataclass
class TaskResult:
    success: bool
    message: str
    data: Any = None
    failed_items: List[str] = field(default_factory=list)


@dataclass
class FileInfo:
    path: str
    name: str
    extension: str
    size: int
    status: FileStatus = FileStatus.PENDING
    chunks: int = 0
    error: str = ""
    
    def to_dict(self) -> Dict:
        return {
            "name": self.name,
            "extension": self.extension,
            "size": self.size,
            "status": self.status.value,
            "chunks": self.chunks,
            "error": self.error
        }


# ============================================================================
# ë¡œê¹… ì„¤ì •
# ============================================================================
def get_app_directory() -> str:
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))


def setup_logger() -> logging.Logger:
    logger = logging.getLogger('RegSearchServer')
    if logger.handlers:
        return logger
    logger.setLevel(logging.INFO)
    log_dir = os.path.join(get_app_directory(), 'logs')
    os.makedirs(log_dir, exist_ok=True)
    fh = logging.FileHandler(
        os.path.join(log_dir, f'server_{datetime.now():%Y%m%d}.log'),
        encoding='utf-8'
    )
    fh.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))
    logger.addHandler(fh)
    ch = logging.StreamHandler()
    ch.setLevel(logging.INFO)
    ch.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))
    logger.addHandler(ch)
    return logger


logger = setup_logger()


# ============================================================================
# ìœ í‹¸ë¦¬í‹°
# ============================================================================
class FileUtils:
    @staticmethod
    def safe_read(path: str, encoding: str = 'utf-8') -> Tuple[Optional[str], Optional[str]]:
        try:
            with open(path, 'r', encoding=encoding, errors='ignore') as f:
                return f.read(), None
        except Exception as e:
            return None, str(e)
    
    @staticmethod
    def get_metadata(path: str) -> Optional[Dict]:
        try:
            stat = os.stat(path)
            return {'size': stat.st_size, 'mtime': stat.st_mtime}
        except OSError as e:
            logger.debug(f"íŒŒì¼ ë©”íƒ€ë°ì´í„° ì¡°íšŒ ì‹¤íŒ¨: {path} - {e}")
            return None
    
    @staticmethod
    def format_size(size: int) -> str:
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size < 1024:
                return f"{size:.1f}{unit}"
            size /= 1024
        return f"{size:.1f}TB"
    
    @staticmethod
    def allowed_file(filename: str) -> bool:
        ext = os.path.splitext(filename)[1].lower()
        return ext in AppConfig.SUPPORTED_EXTENSIONS


# ============================================================================
# BM25 ê²½ëŸ‰ êµ¬í˜„ (ìŠ¤ë ˆë“œ ì•ˆì „)
# ============================================================================
class BM25Light:
    __slots__ = ['k1', 'b', 'corpus', 'doc_lens', 'avgdl', 'idf', 'N', '_lock']
    
    def __init__(self, k1: float = 1.5, b: float = 0.75):
        self.k1 = k1
        self.b = b
        self.corpus: List[List[str]] = []
        self.doc_lens: List[int] = []
        self.avgdl = 0.0
        self.idf: Dict[str, float] = {}
        self.N = 0
        self._lock = threading.RLock()
    
    def _tokenize(self, text: str) -> List[str]:
        if not text:
            return []
        text = re.sub(r'[^\w\sê°€-í£]', ' ', text.lower())
        return [t for t in text.split() if len(t) >= 2]
    
    def fit(self, docs: List[str]):
        with self._lock:
            self.corpus = []
            self.doc_lens = []
            df = Counter()
            for doc in docs:
                tokens = self._tokenize(doc)
                self.corpus.append(tokens)
                self.doc_lens.append(len(tokens))
                df.update(set(tokens))
            self.N = len(docs)
            self.avgdl = sum(self.doc_lens) / self.N if self.N else 0
            self.idf = {t: math.log((self.N - f + 0.5) / (f + 0.5) + 1) for t, f in df.items()}
            del df
            gc.collect()
    
    def search(self, query: str, top_k: int = 5) -> List[Tuple[int, float]]:
        with self._lock:
            if not self.corpus or not query:
                return []
            q_tokens = self._tokenize(query)
            if not q_tokens:
                return []
            scores = []
            for idx, doc_tokens in enumerate(self.corpus):
                if not doc_tokens:
                    continue
                score = self._score(q_tokens, doc_tokens, self.doc_lens[idx])
                if score > 0:
                    scores.append((idx, score))
            scores.sort(key=lambda x: x[1], reverse=True)
            return scores[:top_k]
    
    def _score(self, query: List[str], doc: List[str], doc_len: int) -> float:
        score = 0.0
        doc_tf = Counter(doc)
        for term in query:
            if term not in self.idf:
                continue
            tf = doc_tf.get(term, 0)
            idf = self.idf[term]
            num = tf * (self.k1 + 1)
            den = tf + self.k1 * (1 - self.b + self.b * doc_len / self.avgdl)
            score += idf * num / den if den > 0 else 0
        return score
    
    def clear(self):
        with self._lock:
            self.corpus.clear()
            self.doc_lens.clear()
            self.idf.clear()
            gc.collect()


# ============================================================================
# ë¬¸ì„œ ì¶”ì¶œê¸°
# ============================================================================
class DocumentExtractor:
    def __init__(self):
        self._docx_module = None
        self._pdf_module = None
    
    @property
    def docx(self):
        if self._docx_module is None:
            try:
                from docx import Document
                self._docx_module = Document
            except ImportError:
                self._docx_module = False
        return self._docx_module
    
    @property
    def pdf(self):
        if self._pdf_module is None:
            try:
                from pypdf import PdfReader
                self._pdf_module = PdfReader
            except ImportError:
                self._pdf_module = False
        return self._pdf_module
    
    def extract(self, path: str) -> Tuple[str, Optional[str]]:
        if not path or not os.path.exists(path):
            return "", f"íŒŒì¼ ì—†ìŒ: {path}"
        if not os.path.isfile(path):
            return "", f"íŒŒì¼ì´ ì•„ë‹˜: {path}"
        ext = os.path.splitext(path)[1].lower()
        if ext == '.txt':
            return self._extract_txt(path)
        elif ext == '.docx':
            return self._extract_docx(path)
        elif ext == '.pdf':
            return self._extract_pdf(path)
        return "", f"ì§€ì›í•˜ì§€ ì•ŠëŠ” í˜•ì‹: {ext}"
    
    def _extract_txt(self, path: str) -> Tuple[str, Optional[str]]:
        return FileUtils.safe_read(path)
    
    def _extract_docx(self, path: str) -> Tuple[str, Optional[str]]:
        if not self.docx:
            return "", "DOCX ë¼ì´ë¸ŒëŸ¬ë¦¬ ì—†ìŒ (pip install python-docx)"
        try:
            doc = self.docx(path)
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(' | '.join(cells))
            return '\n\n'.join(parts), None
        except Exception as e:
            return "", f"DOCX ì˜¤ë¥˜: {e}"
    
    def _extract_pdf(self, path: str) -> Tuple[str, Optional[str]]:
        if not self.pdf:
            return "", "PDF ë¼ì´ë¸ŒëŸ¬ë¦¬ ì—†ìŒ (pip install pypdf)"
        try:
            reader = self.pdf(path)
            if reader.is_encrypted:
                try:
                    reader.decrypt('')
                except Exception:
                    return "", "ì•”í˜¸í™”ëœ PDF"
            texts = []
            for page in reader.pages:
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        texts.append(text.strip())
                except Exception:
                    continue
            if not texts:
                return "", "í…ìŠ¤íŠ¸ ì—†ìŒ (ì´ë¯¸ì§€ PDF)"
            return '\n\n'.join(texts), None
        except Exception as e:
            return "", f"PDF ì˜¤ë¥˜: {e}"


# ============================================================================
# ê²€ìƒ‰ ìºì‹œ (LRU)
# ============================================================================
class SearchCache:
    def __init__(self, max_size: int = 100):
        self.cache: Dict[str, Tuple[float, Any]] = {}
        self.max_size = max_size
        self._lock = threading.Lock()
    
    def _make_key(self, query: str, k: int, hybrid: bool) -> str:
        return f"{query}|{k}|{hybrid}"
    
    def get(self, query: str, k: int, hybrid: bool) -> Optional[Any]:
        key = self._make_key(query, k, hybrid)
        with self._lock:
            if key in self.cache:
                timestamp, result = self.cache[key]
                # 5ë¶„ ì´ë‚´ ìºì‹œë§Œ ìœ íš¨
                if time.time() - timestamp < 300:
                    return result
                del self.cache[key]
        return None
    
    def set(self, query: str, k: int, hybrid: bool, result: Any):
        key = self._make_key(query, k, hybrid)
        with self._lock:
            if len(self.cache) >= self.max_size:
                # ê°€ì¥ ì˜¤ë˜ëœ í•­ëª© ì œê±°
                oldest_key = min(self.cache.keys(), key=lambda x: self.cache[x][0])
                del self.cache[oldest_key]
            self.cache[key] = (time.time(), result)
    
    def clear(self):
        with self._lock:
            self.cache.clear()


# ============================================================================
# Rate Limiter (IP ê¸°ë°˜ ìš”ì²­ ì œí•œ)
# ============================================================================
class RateLimiter:
    """IP ê¸°ë°˜ ìš”ì²­ ì œí•œìœ¼ë¡œ ì„œë²„ ê³¼ë¶€í•˜ ë°©ì§€"""
    
    def __init__(self, requests_per_minute: int = 60):
        self.requests: Dict[str, List[float]] = {}  # ip -> [timestamps]
        self.limit = requests_per_minute
        self._lock = threading.Lock()
        self._cleanup_interval = 60  # ì •ë¦¬ ì£¼ê¸° (ì´ˆ)
        self._last_cleanup = time.time()
    
    def is_allowed(self, ip: str) -> bool:
        """ìš”ì²­ì´ í—ˆìš©ë˜ëŠ”ì§€ í™•ì¸"""
        current_time = time.time()
        
        with self._lock:
            # ì£¼ê¸°ì  ì •ë¦¬
            if current_time - self._last_cleanup > self._cleanup_interval:
                self._cleanup(current_time)
                self._last_cleanup = current_time
            
            if ip not in self.requests:
                self.requests[ip] = []
            
            # 1ë¶„ ì´ë‚´ ìš”ì²­ë§Œ ìœ ì§€
            cutoff = current_time - 60
            self.requests[ip] = [t for t in self.requests[ip] if t > cutoff]
            
            # ì œí•œ í™•ì¸
            if len(self.requests[ip]) >= self.limit:
                logger.warning(f"Rate limit exceeded for IP: {ip}")
                return False
            
            # ìƒˆ ìš”ì²­ ê¸°ë¡
            self.requests[ip].append(current_time)
            return True
    
    def _cleanup(self, current_time: float):
        """ì˜¤ë˜ëœ ê¸°ë¡ ì •ë¦¬"""
        cutoff = current_time - 60
        expired_ips = []
        for ip, timestamps in self.requests.items():
            self.requests[ip] = [t for t in timestamps if t > cutoff]
            if not self.requests[ip]:
                expired_ips.append(ip)
        for ip in expired_ips:
            del self.requests[ip]
    
    def get_remaining(self, ip: str) -> int:
        """ë‚¨ì€ ìš”ì²­ ìˆ˜ ë°˜í™˜"""
        with self._lock:
            current_time = time.time()
            cutoff = current_time - 60
            if ip not in self.requests:
                return self.limit
            recent = [t for t in self.requests[ip] if t > cutoff]
            return max(0, self.limit - len(recent))


# ============================================================================
# ê²€ìƒ‰ ìš”ì²­ í (ë™ì‹œ ê²€ìƒ‰ ìˆ˜ ì œí•œ)
# ============================================================================
class SearchQueue:
    """ë™ì‹œ ê²€ìƒ‰ ìˆ˜ë¥¼ ì œí•œí•˜ì—¬ ì„œë²„ ì•ˆì •ì„± í™•ë³´"""
    
    def __init__(self, max_concurrent: int = 10):
        self._semaphore = threading.Semaphore(max_concurrent)
        self._active_count = 0
        self._lock = threading.Lock()
        self._total_processed = 0
        self._total_rejected = 0
    
    def acquire(self, timeout: float = 30.0) -> bool:
        """ê²€ìƒ‰ ìŠ¬ë¡¯ íšë“ (íƒ€ì„ì•„ì›ƒ ì§€ì›)"""
        acquired = self._semaphore.acquire(timeout=timeout)
        if acquired:
            with self._lock:
                self._active_count += 1
        else:
            with self._lock:
                self._total_rejected += 1
            logger.warning("Search queue full, request rejected")
        return acquired
    
    def release(self):
        """ê²€ìƒ‰ ìŠ¬ë¡¯ ë°˜í™˜"""
        self._semaphore.release()
        with self._lock:
            self._active_count = max(0, self._active_count - 1)
            self._total_processed += 1
    
    def get_stats(self) -> Dict:
        """í ìƒíƒœ ë°˜í™˜"""
        with self._lock:
            return {
                'active': self._active_count,
                'processed': self._total_processed,
                'rejected': self._total_rejected
            }


# ì „ì—­ Rate Limiter ë° Search Queue ì¸ìŠ¤í„´ìŠ¤
rate_limiter = RateLimiter(AppConfig.RATE_LIMIT_PER_MINUTE)
search_queue = SearchQueue(AppConfig.MAX_CONCURRENT_SEARCHES)

# íŒŒì¼ ì‘ì—… ë½ (ë™ì‹œ íŒŒì¼ ì—…ë¡œë“œ/ì‚­ì œ ë³´í˜¸)
file_operation_lock = threading.Lock()

# ============================================================================
# ê²€ìƒ‰ íˆìŠ¤í† ë¦¬ (ìµœê·¼ ê²€ìƒ‰ì–´ + ì¸ê¸° ê²€ìƒ‰ì–´)
# ============================================================================
class SearchHistory:
    def __init__(self, max_recent: int = 20, max_popular: int = 10):
        self.recent: List[Dict] = []  # [{query, timestamp}, ...]
        self.popular: Counter = Counter()  # query -> count
        self.max_recent = max_recent
        self.max_popular = max_popular
        self._lock = threading.Lock()
    
    def add(self, query: str):
        """ê²€ìƒ‰ì–´ ì¶”ê°€"""
        query = query.strip()
        if len(query) < 2:
            return
        
        with self._lock:
            # ìµœê·¼ ê²€ìƒ‰ì–´ì—ì„œ ì¤‘ë³µ ì œê±°
            self.recent = [r for r in self.recent if r['query'] != query]
            # ìƒˆ ê²€ìƒ‰ì–´ ì¶”ê°€
            self.recent.insert(0, {
                'query': query,
                'timestamp': time.time()
            })
            # ìµœëŒ€ í¬ê¸° ìœ ì§€
            self.recent = self.recent[:self.max_recent]
            # ì¸ê¸° ê²€ìƒ‰ì–´ ì—…ë°ì´íŠ¸
            self.popular[query] += 1
    
    def get_recent(self, limit: int = 10) -> List[str]:
        """ìµœê·¼ ê²€ìƒ‰ì–´ ë°˜í™˜"""
        with self._lock:
            return [r['query'] for r in self.recent[:limit]]
    
    def get_popular(self, limit: int = 10) -> List[Tuple[str, int]]:
        """ì¸ê¸° ê²€ìƒ‰ì–´ ë°˜í™˜ (ê²€ìƒ‰ì–´, íšŸìˆ˜)"""
        with self._lock:
            return self.popular.most_common(min(limit, self.max_popular))
    
    def suggest(self, prefix: str, limit: int = 5) -> List[str]:
        """ì ‘ë‘ì‚¬ ê¸°ë°˜ ê²€ìƒ‰ì–´ ì¶”ì²œ"""
        prefix = prefix.strip().lower()
        if len(prefix) < 1:
            return []
        
        with self._lock:
            suggestions = []
            # ìµœê·¼ ê²€ìƒ‰ì–´ì—ì„œ ë§¤ì¹­
            for r in self.recent:
                if r['query'].lower().startswith(prefix):
                    suggestions.append(r['query'])
            # ì¸ê¸° ê²€ìƒ‰ì–´ì—ì„œ ë§¤ì¹­
            for q, _ in self.popular.most_common():
                if q.lower().startswith(prefix) and q not in suggestions:
                    suggestions.append(q)
            return suggestions[:limit]
    
    def clear(self):
        """íˆìŠ¤í† ë¦¬ ì´ˆê¸°í™”"""
        with self._lock:
            self.recent.clear()
            self.popular.clear()


# ============================================================================
# í…ìŠ¤íŠ¸ í•˜ì´ë¼ì´í„°
# ============================================================================
class TextHighlighter:
    @staticmethod
    def highlight(text: str, query: str, tag: str = 'mark') -> str:
        """ê²€ìƒ‰ì–´ë¥¼ íƒœê·¸ë¡œ ê°ì‹¸ì„œ í•˜ì´ë¼ì´íŠ¸"""
        if not text or not query:
            return text
        
        # ê²€ìƒ‰ì–´ë¥¼ ê³µë°±ìœ¼ë¡œ ë¶„ë¦¬
        keywords = [kw.strip() for kw in query.split() if len(kw.strip()) >= 2]
        if not keywords:
            return text
        
        # ê° í‚¤ì›Œë“œì— ëŒ€í•´ í•˜ì´ë¼ì´íŠ¸ ì ìš©
        result = text
        for keyword in keywords:
            # ëŒ€ì†Œë¬¸ì ë¬´ì‹œ ê²€ìƒ‰
            pattern = re.compile(re.escape(keyword), re.IGNORECASE)
            result = pattern.sub(f'<{tag}>\\g<0></{tag}>', result)
        
        return result
    
    @staticmethod
    def extract_keywords(documents: List[str], top_k: int = 50) -> List[str]:
        """ë¬¸ì„œì—ì„œ í•µì‹¬ í‚¤ì›Œë“œ ì¶”ì¶œ"""
        if not documents:
            return []
        
        # ê°„ë‹¨í•œ í‚¤ì›Œë“œ ì¶”ì¶œ (ë¹ˆë„ ê¸°ë°˜)
        word_freq = Counter()
        for doc in documents:
            # í•œê¸€, ì˜ë¬¸ ë‹¨ì–´ ì¶”ì¶œ
            words = re.findall(r'[ê°€-í£]{2,}|[a-zA-Z]{3,}', doc)
            word_freq.update(words)
        
        # ë¶ˆìš©ì–´ í•„í„°ë§ (ê°„ë‹¨í•œ í•œê¸€ ë¶ˆìš©ì–´)
        stopwords = {'ìˆëŠ”', 'í•˜ëŠ”', 'ë°', 'ë“±', 'ì´', 'ê°€', 'ì„', 'ë¥¼', 'ì˜', 'ì—', 'ë¡œ', 'ìœ¼ë¡œ'}
        keywords = [w for w, _ in word_freq.most_common(top_k * 2) if w not in stopwords]
        
        return keywords[:top_k]


# ============================================================================
# í•µì‹¬ QA ì‹œìŠ¤í…œ (ìŠ¤ë ˆë“œ ì•ˆì „)
# ============================================================================
class RegulationQASystem:
    def __init__(self):
        self.vector_store = None
        self.embedding_model = None
        self.model_id = None
        self.model_name = ""
        self.extractor = DocumentExtractor()
        self.cache_path = os.path.join(tempfile.gettempdir(), "reg_qa_server_v10")
        self.bm25 = None
        self.documents: List[str] = []
        self.doc_meta: List[Dict] = []
        self.file_infos: Dict[str, FileInfo] = {}
        self.current_folder = ""
        self._lock = threading.RLock()
        self._search_cache = SearchCache(AppConfig.SEARCH_CACHE_SIZE)
        self._search_history = SearchHistory()  # ê²€ìƒ‰ íˆìŠ¤í† ë¦¬
        self._keyword_cache: List[str] = []  # ë¬¸ì„œ í‚¤ì›Œë“œ ìºì‹œ
        self._executor = ThreadPoolExecutor(max_workers=AppConfig.MAX_WORKERS)
        self._is_ready = False
        self._is_loading = False
        self._load_progress = ""
        self._load_error = ""  # ë§ˆì§€ë§‰ ë¡œë“œ ì˜¤ë¥˜ ë©”ì‹œì§€
    
    def get_keywords(self, limit: int = 50) -> List[str]:
        """ë¬¸ì„œì—ì„œ ì¶”ì¶œí•œ í‚¤ì›Œë“œ ë°˜í™˜ (ìë™ì™„ì„±ìš©)"""
        if not self._keyword_cache and self.documents:
            self._keyword_cache = TextHighlighter.extract_keywords(self.documents, limit)
        return self._keyword_cache[:limit]
    
    @property
    def is_ready(self) -> bool:
        return self._is_ready and self.embedding_model is not None
    
    @property
    def is_loading(self) -> bool:
        return self._is_loading
    
    @property
    def load_progress(self) -> str:
        return self._load_progress
    
    @property
    def load_error(self) -> str:
        """ë§ˆì§€ë§‰ ë¡œë“œ ì˜¤ë¥˜ ë©”ì‹œì§€"""
        return self._load_error
    
    def load_model(self, model_name: str, offline_mode: bool = None, local_model_path: str = None) -> TaskResult:
        """AI ì„ë² ë”© ëª¨ë¸ ë¡œë“œ
        
        Args:
            model_name: ëª¨ë¸ ì´ë¦„ (AVAILABLE_MODELSì˜ í‚¤)
            offline_mode: ì˜¤í”„ë¼ì¸ ëª¨ë“œ ì—¬ë¶€ (Noneì´ë©´ AppConfig ì„¤ì • ì‚¬ìš©)
            local_model_path: ë¡œì»¬ ëª¨ë¸ ê²½ë¡œ (Noneì´ë©´ AppConfig ì„¤ì • ì‚¬ìš©)
        """
        import traceback
        
        if self._is_loading:
            return TaskResult(False, "ì´ë¯¸ ëª¨ë¸ì„ ë¡œë”© ì¤‘ì…ë‹ˆë‹¤")
        
        # ì˜¤í”„ë¼ì¸ ëª¨ë“œ ì„¤ì • ê²°ì •
        is_offline = offline_mode if offline_mode is not None else AppConfig.OFFLINE_MODE
        model_path_override = local_model_path if local_model_path is not None else AppConfig.LOCAL_MODEL_PATH
        
        model_id = AppConfig.AVAILABLE_MODELS.get(model_name, AppConfig.AVAILABLE_MODELS[AppConfig.DEFAULT_MODEL])
        self._load_error = ""  # ì˜¤ë¥˜ ì´ˆê¸°í™”
        
        try:
            self._is_loading = True
            self._load_progress = "ë¼ì´ë¸ŒëŸ¬ë¦¬ ë¡œë“œ ì¤‘..."
            logger.info("ë¼ì´ë¸ŒëŸ¬ë¦¬ ë¡œë“œ ì¤‘...")
            
            # PyTorch ë¡œë“œ ì‹œë„
            try:
                import torch
                logger.info(f"PyTorch ë²„ì „: {torch.__version__}")
            except ImportError as e:
                error_msg = f"PyTorch ë¡œë“œ ì‹¤íŒ¨: {e}"
                logger.error(error_msg)
                self._load_error = error_msg
                return TaskResult(False, error_msg)
            
            # LangChain HuggingFace ë¡œë“œ ì‹œë„
            try:
                from langchain_huggingface import HuggingFaceEmbeddings
                logger.info("LangChain HuggingFace ë¡œë“œ ì™„ë£Œ")
            except ImportError as e:
                error_msg = f"LangChain HuggingFace ë¡œë“œ ì‹¤íŒ¨: {e}"
                logger.error(error_msg)
                self._load_error = error_msg
                return TaskResult(False, error_msg)
            
            # ëª¨ë¸ ìºì‹œ ê²½ë¡œ ê²°ì •
            if model_path_override:
                cache_dir = model_path_override
            else:
                cache_dir = os.path.join(get_app_directory(), 'models')
            os.makedirs(cache_dir, exist_ok=True)
            
            # ì˜¤í”„ë¼ì¸ ëª¨ë“œ ë¡œê¹… ë° í™˜ê²½ë³€ìˆ˜ ì„¤ì •
            if is_offline:
                self._load_progress = "ì˜¤í”„ë¼ì¸ ëª¨ë“œ: ë¡œì»¬ ëª¨ë¸ ë¡œë”© ì¤‘..."
                logger.info(f"ğŸ”’ ì˜¤í”„ë¼ì¸ ëª¨ë“œ í™œì„±í™”: ë¡œì»¬ ëª¨ë¸ë§Œ ì‚¬ìš©")
                logger.info(f"ğŸ“‚ ëª¨ë¸ ê²½ë¡œ: {cache_dir}")
                
                # HuggingFace ì˜¤í”„ë¼ì¸ í™˜ê²½ë³€ìˆ˜ ì„¤ì • (ë„¤íŠ¸ì›Œí¬ ìš”ì²­ ë°©ì§€)
                os.environ['HF_HUB_OFFLINE'] = '1'
                os.environ['TRANSFORMERS_OFFLINE'] = '1'
                os.environ['HF_DATASETS_OFFLINE'] = '1'
                
                # ëª¨ë¸ ë””ë ‰í† ë¦¬ ì¡´ì¬ í™•ì¸
                model_subdir = os.path.join(cache_dir, model_id.replace('/', '--'))
                if not os.path.exists(model_subdir):
                    # HuggingFace ìºì‹œ í˜•ì‹ë„ í™•ì¸
                    alt_dirs = [
                        os.path.join(cache_dir, 'models--' + model_id.replace('/', '--')),
                        cache_dir  # ì§ì ‘ ê²½ë¡œ
                    ]
                    found = False
                    for alt_dir in alt_dirs:
                        if os.path.exists(alt_dir):
                            found = True
                            break
                    
                    if not found:
                        error_msg = (
                            f"ì˜¤í”„ë¼ì¸ ëª¨ë“œì—ì„œ ëª¨ë¸ì„ ì°¾ì„ ìˆ˜ ì—†ìŠµë‹ˆë‹¤.\n"
                            f"ëª¨ë¸ ê²½ë¡œ: {cache_dir}\n"
                            f"ì˜ˆìƒ ëª¨ë¸ í´ë”: {model_id.replace('/', '--')}\n\n"
                            f"í•´ê²° ë°©ë²•:\n"
                            f"1. ì¸í„°ë„· ì—°ê²°ëœ í™˜ê²½ì—ì„œ 'python download_models.py' ì‹¤í–‰\n"
                            f"2. ìƒì„±ëœ models í´ë”ë¥¼ ì´ ì„œë²„ë¡œ ë³µì‚¬\n"
                            f"3. ì„¤ì •ì—ì„œ 'ë¡œì»¬ ëª¨ë¸ ê²½ë¡œ'ë¥¼ ì˜¬ë°”ë¥´ê²Œ ì§€ì •"
                        )
                        logger.error(error_msg)
                        self._load_error = error_msg
                        return TaskResult(False, error_msg)
            else:
                self._load_progress = "ëª¨ë¸ ë‹¤ìš´ë¡œë“œ/ë¡œë”© ì¤‘..."
                logger.info(f"ëª¨ë¸ ë¡œë”© ì¤‘: {model_name}")
                logger.info(f"ëª¨ë¸ ìºì‹œ ê²½ë¡œ: {cache_dir}")
            
            device = 'cuda' if torch.cuda.is_available() else 'cpu'
            logger.info(f"ì‚¬ìš© ë””ë°”ì´ìŠ¤: {device}")
            
            # ëª¨ë¸ ë¡œë“œ ì˜µì…˜ ì„¤ì •
            model_kwargs = {'device': device}
            if is_offline:
                model_kwargs['local_files_only'] = True
            
            self.embedding_model = HuggingFaceEmbeddings(
                model_name=model_id,
                cache_folder=cache_dir,
                model_kwargs=model_kwargs,
                encode_kwargs={'normalize_embeddings': True}
            )
            self.model_id = model_id
            self.model_name = model_name
            self._is_ready = True
            
            gc.collect()
            if device == 'cuda':
                torch.cuda.empty_cache()
            
            mode_str = "ì˜¤í”„ë¼ì¸" if is_offline else "ì˜¨ë¼ì¸"
            self._load_progress = "ì™„ë£Œ"
            logger.info(f"ëª¨ë¸ ë¡œë“œ ì™„ë£Œ: {model_name} ({device}, {mode_str})")
            return TaskResult(True, f"ëª¨ë¸ ë¡œë“œ ì™„ë£Œ ({device}, {mode_str})")
            
        except Exception as e:
            error_detail = traceback.format_exc()
            logger.error(f"ëª¨ë¸ ë¡œë“œ ì‹¤íŒ¨: {e}\n{error_detail}")
            self._load_error = str(e)
            self._load_progress = f"ì‹¤íŒ¨: {e}"
            return TaskResult(False, f"ëª¨ë¸ ë¡œë“œ ì‹¤íŒ¨: {e}")
        finally:
            self._is_loading = False
    
    def _get_cache_dir(self, folder: str) -> str:
        if not self.model_id:
            raise ValueError("ëª¨ë¸ì´ ë¡œë“œë˜ì§€ ì•Šì•˜ìŠµë‹ˆë‹¤")
        h1 = hashlib.md5(self.model_id.encode()).hexdigest()[:6]
        h2 = hashlib.md5(folder.encode()).hexdigest()[:6]
        return os.path.join(self.cache_path, f"{h2}_{h1}")
    
    def process_documents(self, folder: str, files: List[str], progress_cb=None) -> TaskResult:
        """ë¬¸ì„œ ì²˜ë¦¬ ë° ì¸ë±ì‹±"""
        if not self.embedding_model:
            return TaskResult(False, "ëª¨ë¸ì´ ë¡œë“œë˜ì§€ ì•Šì•˜ìŠµë‹ˆë‹¤")
        
        with self._lock:
            return self._process_internal(folder, files, progress_cb)
    
    def _process_internal(self, folder: str, files: List[str], progress_cb) -> TaskResult:
        # LangChain ìµœì‹  ë²„ì „ í˜¸í™˜ import
        try:
            from langchain_text_splitters import CharacterTextSplitter
        except ImportError:
            from langchain.text_splitter import CharacterTextSplitter
        
        from langchain_community.vectorstores import FAISS
        
        try:
            from langchain_core.documents import Document
        except ImportError:
            from langchain.docstore.document import Document
        
        self.current_folder = folder
        cache_dir = self._get_cache_dir(folder)
        self.file_infos.clear()
        self._search_cache.clear()
        
        # íŒŒì¼ ì •ë³´ ì´ˆê¸°í™”
        for fp in files:
            meta = FileUtils.get_metadata(fp)
            self.file_infos[fp] = FileInfo(
                fp, os.path.basename(fp),
                os.path.splitext(fp)[1].lower(),
                meta['size'] if meta else 0
            )
        
        if progress_cb:
            progress_cb(5, "ìºì‹œ í™•ì¸...")
        
        cache_info = self._load_cache_info(cache_dir)
        to_process, cached = [], []
        
        for fp in files:
            fname = os.path.basename(fp)
            meta = FileUtils.get_metadata(fp)
            if meta and fname in cache_info:
                cm = cache_info[fname]
                if cm.get('size') == meta['size'] and cm.get('mtime') == meta['mtime']:
                    cached.append(fp)
                    self.file_infos[fp].status = FileStatus.CACHED
                    self.file_infos[fp].chunks = cm.get('chunks', 0)
                    continue
            to_process.append(fp)
        
        self.documents, self.doc_meta = [], []
        
        # ìºì‹œëœ ë°ì´í„° ë¡œë“œ
        if cached and os.path.exists(os.path.join(cache_dir, "index.faiss")):
            try:
                if progress_cb:
                    progress_cb(10, "ìºì‹œ ë¡œë“œ...")
                self.vector_store = FAISS.load_local(
                    cache_dir, self.embedding_model,
                    allow_dangerous_deserialization=True
                )
                docs_path = os.path.join(cache_dir, "docs.json")
                if os.path.exists(docs_path):
                    with open(docs_path, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                        self.documents = data.get('docs', [])
                        self.doc_meta = data.get('meta', [])
            except Exception as e:
                logger.warning(f"ìºì‹œ ë¡œë“œ ì‹¤íŒ¨: {e}")
                to_process, cached = files, []
                self.vector_store = None
        
        if not to_process:
            self._build_bm25()
            if progress_cb:
                progress_cb(100, "ì™„ë£Œ!")
            return TaskResult(
                True,
                f"ìºì‹œì—ì„œ {len(cached)}ê°œ íŒŒì¼ ë¡œë“œ",
                {'chunks': len(self.documents), 'cached': len(cached), 'new': 0}
            )
        
        splitter = CharacterTextSplitter(
            separator="\n\n",
            chunk_size=AppConfig.CHUNK_SIZE,
            chunk_overlap=AppConfig.CHUNK_OVERLAP
        )
        failed, new_docs, new_cache_info = [], [], {}
        
        # ë³‘ë ¬ ë¬¸ì„œ ì¶”ì¶œ í•¨ìˆ˜
        def extract_file(fp: str) -> Tuple[str, str, Optional[str], Optional[Dict]]:
            """íŒŒì¼ì—ì„œ í…ìŠ¤íŠ¸ ì¶”ì¶œ (ë³‘ë ¬ ì²˜ë¦¬ìš©)"""
            fname = os.path.basename(fp)
            try:
                content, error = self.extractor.extract(fp)
                meta = FileUtils.get_metadata(fp)
                return fp, fname, content, error, meta
            except Exception as e:
                return fp, fname, None, str(e), None
        
        # ThreadPoolExecutorë¡œ ë³‘ë ¬ ë¬¸ì„œ ì¶”ì¶œ
        extracted_results = []
        if progress_cb:
            progress_cb(15, f"ë¬¸ì„œ ì¶”ì¶œ ì¤‘... (ë³‘ë ¬ ì²˜ë¦¬)")
        
        with ThreadPoolExecutor(max_workers=min(AppConfig.MAX_WORKERS, len(to_process))) as executor:
            futures = {executor.submit(extract_file, fp): fp for fp in to_process}
            completed = 0
            for future in futures:
                try:
                    result = future.result(timeout=60)  # íŒŒì¼ë‹¹ 60ì´ˆ íƒ€ì„ì•„ì›ƒ
                    extracted_results.append(result)
                    completed += 1
                    if progress_cb and completed % 5 == 0:
                        progress = 15 + int((completed / len(to_process)) * 30)
                        progress_cb(progress, f"ì¶”ì¶œ ì™„ë£Œ: {completed}/{len(to_process)}")
                except Exception as e:
                    fp = futures[future]
                    fname = os.path.basename(fp)
                    extracted_results.append((fp, fname, None, f"ì¶”ì¶œ íƒ€ì„ì•„ì›ƒ: {e}", None))
        
        # ì¶”ì¶œ ê²°ê³¼ ì²˜ë¦¬ ë° ì²­í‚¹
        if progress_cb:
            progress_cb(50, "í…ìŠ¤íŠ¸ ì²­í‚¹ ì¤‘...")
        
        for fp, fname, content, error, meta in extracted_results:
            self.file_infos[fp].status = FileStatus.PROCESSING
            
            if error:
                failed.append(f"{fname} ({error})")
                self.file_infos[fp].status = FileStatus.FAILED
                self.file_infos[fp].error = error
                continue
            if not content or not content.strip():
                failed.append(f"{fname} (ë¹ˆ íŒŒì¼)")
                self.file_infos[fp].status = FileStatus.FAILED
                self.file_infos[fp].error = "ë¹ˆ íŒŒì¼"
                continue
            
            try:
                chunks = splitter.split_text(content)
                chunk_count = 0
                for chunk in chunks:
                    if chunk.strip():
                        new_docs.append(Document(
                            page_content=chunk.strip(),
                            metadata={"source": fname, "path": fp}
                        ))
                        self.documents.append(chunk.strip())
                        self.doc_meta.append({"source": fname, "path": fp})
                        chunk_count += 1
                
                self.file_infos[fp].status = FileStatus.SUCCESS
                self.file_infos[fp].chunks = chunk_count
                
                if meta:
                    new_cache_info[fname] = {
                        'size': meta['size'],
                        'mtime': meta['mtime'],
                        'chunks': chunk_count
                    }
            except Exception as e:
                failed.append(f"{fname} ({e})")
                self.file_infos[fp].status = FileStatus.FAILED
                self.file_infos[fp].error = str(e)

        
        if not new_docs and not self.vector_store:
            return TaskResult(False, "ì²˜ë¦¬ ê°€ëŠ¥í•œ ë¬¸ì„œ ì—†ìŒ", failed_items=failed)
        
        if progress_cb:
            progress_cb(75, "ë²¡í„° ì¸ë±ìŠ¤ ìƒì„±...")
        
        try:
            if new_docs:
                if self.vector_store:
                    batch_size = 100
                    for i in range(0, len(new_docs), batch_size):
                        self.vector_store.add_documents(new_docs[i:i + batch_size])
                else:
                    self.vector_store = FAISS.from_documents(new_docs, self.embedding_model)
        except Exception as e:
            return TaskResult(False, f"ì¸ë±ìŠ¤ ìƒì„± ì‹¤íŒ¨: {e}")
        
        if progress_cb:
            progress_cb(85, "í‚¤ì›Œë“œ ì¸ë±ìŠ¤ ìƒì„±...")
        self._build_bm25()
        
        if progress_cb:
            progress_cb(90, "ìºì‹œ ì €ì¥...")
        self._save_cache(cache_dir, cache_info, new_cache_info)
        
        gc.collect()
        if progress_cb:
            progress_cb(100, "ì™„ë£Œ!")
        
        return TaskResult(
            True,
            f"{len(files) - len(failed)}ê°œ ì²˜ë¦¬ ì™„ë£Œ",
            {
                'chunks': len(self.documents),
                'new': len(to_process) - len(failed),
                'cached': len(cached)
            },
            failed
        )
    
    def _build_bm25(self):
        if self.documents:
            self.bm25 = BM25Light()
            self.bm25.fit(self.documents)
    
    def _load_cache_info(self, cache_dir: str) -> Dict:
        path = os.path.join(cache_dir, "cache_info.json")
        if os.path.exists(path):
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except (json.JSONDecodeError, IOError) as e:
                logger.debug(f"ìºì‹œ ì •ë³´ ë¡œë“œ ì‹¤íŒ¨: {e}")
        return {}
    
    def _save_cache(self, cache_dir: str, old_info: Dict, new_info: Dict):
        try:
            os.makedirs(cache_dir, exist_ok=True)
            self.vector_store.save_local(cache_dir)
            with open(os.path.join(cache_dir, "cache_info.json"), 'w', encoding='utf-8') as f:
                json.dump({**old_info, **new_info}, f, ensure_ascii=False)
            with open(os.path.join(cache_dir, "docs.json"), 'w', encoding='utf-8') as f:
                json.dump({'docs': self.documents, 'meta': self.doc_meta}, f, ensure_ascii=False)
        except Exception as e:
            logger.warning(f"ìºì‹œ ì €ì¥ ì‹¤íŒ¨: {e}")
    
    def search(self, query: str, k: int = 3, hybrid: bool = True, 
                filter_file: str = None, sort_by: str = 'relevance') -> TaskResult:
        """í•˜ì´ë¸Œë¦¬ë“œ ê²€ìƒ‰ ìˆ˜í–‰
        
        Args:
            query: ê²€ìƒ‰ì–´
            k: ê²°ê³¼ ê°œìˆ˜
            hybrid: í•˜ì´ë¸Œë¦¬ë“œ ê²€ìƒ‰ ì‚¬ìš© ì—¬ë¶€
            filter_file: íŠ¹ì • íŒŒì¼ì—ì„œë§Œ ê²€ìƒ‰ (None=ì „ì²´)
            sort_by: ì •ë ¬ ë°©ì‹ ('relevance', 'filename', 'length')
        """
        if not self.vector_store:
            return TaskResult(False, "ë¬¸ì„œê°€ ë¡œë“œë˜ì§€ ì•ŠìŒ")
        
        query = query.strip()
        if len(query) < 2:
            return TaskResult(False, "ê²€ìƒ‰ì–´ê°€ ë„ˆë¬´ ì§§ìŠµë‹ˆë‹¤ (ìµœì†Œ 2ì)")
        
        # ìºì‹œ í™•ì¸
        cached_result = self._search_cache.get(query, k, hybrid)
        if cached_result is not None:
            return TaskResult(True, "ê²€ìƒ‰ ì™„ë£Œ (ìºì‹œ)", cached_result)
        
        try:
            k = max(1, min(k, AppConfig.MAX_SEARCH_RESULTS))
            
            # ë²¡í„° ê²€ìƒ‰
            vec_results = self.vector_store.similarity_search_with_score(query, k=k*2)
            
            results = {}
            if vec_results and len(vec_results) > 0:
                distances = [r[1] for r in vec_results]
                min_d = min(distances)
                max_d = max(distances)
                rng = max_d - min_d if max_d != min_d else 1
                
                for doc, dist in vec_results:
                    key = doc.page_content[:100]
                    score = max(0.1, 1 - ((dist - min_d) / (rng + 0.001)))
                    results[key] = {
                        'content': doc.page_content,
                        'source': doc.metadata.get('source', '?'),
                        'path': doc.metadata.get('path', ''),
                        'vec_score': score,
                        'bm25_score': 0
                    }
            
            # BM25 ê²€ìƒ‰ (í•˜ì´ë¸Œë¦¬ë“œ)
            if hybrid and self.bm25:
                try:
                    bm_res = self.bm25.search(query, top_k=k*2)
                except Exception as bm_err:
                    logger.warning(f"BM25 ê²€ìƒ‰ ì˜¤ë¥˜: {bm_err}")
                    bm_res = []
                if bm_res and len(bm_res) > 0:
                    bm_scores = [r[1] for r in bm_res]
                    max_bm = max(bm_scores) if bm_scores else 1
                    for idx, sc in bm_res:
                        if 0 <= idx < len(self.documents):
                            key = self.documents[idx][:100]
                            norm = sc / (max_bm + 0.001)
                            if key in results:
                                results[key]['bm25_score'] = norm
                            else:
                                meta = self.doc_meta[idx] if idx < len(self.doc_meta) else {}
                                results[key] = {
                                    'content': self.documents[idx],
                                    'source': meta.get('source', '?'),
                                    'path': meta.get('path', ''),
                                    'vec_score': 0,
                                    'bm25_score': norm
                                }
            
            # íŒŒì¼ í•„í„°ë§ ì ìš©
            if filter_file:
                results = {k: v for k, v in results.items() if v['source'] == filter_file}
            
            # ìµœì¢… ì ìˆ˜ ê³„ì‚°
            for item in results.values():
                item['score'] = (
                    AppConfig.VECTOR_WEIGHT * item['vec_score'] +
                    AppConfig.BM25_WEIGHT * item['bm25_score']
                )
            
            # ì •ë ¬ ì ìš©
            if sort_by == 'filename':
                sorted_res = sorted(results.values(), key=lambda x: x['source'])[:k]
            elif sort_by == 'length':
                sorted_res = sorted(results.values(), key=lambda x: len(x['content']), reverse=True)[:k]
            else:  # relevance (ê¸°ë³¸)
                sorted_res = sorted(results.values(), key=lambda x: x['score'], reverse=True)[:k]
            
            # ìºì‹œ ì €ì¥ (í•„í„° ì—†ì„ ë•Œë§Œ)
            if not filter_file:
                self._search_cache.set(query, k, hybrid, sorted_res)
            
            return TaskResult(True, "ê²€ìƒ‰ ì™„ë£Œ", sorted_res)
            
        except Exception as e:
            logger.error(f"ê²€ìƒ‰ ì˜¤ë¥˜: {e}")
            return TaskResult(False, f"ê²€ìƒ‰ ì˜¤ë¥˜: {e}")
    
    def get_file_infos(self) -> List[Dict]:
        return [info.to_dict() for info in self.file_infos.values()]
    
    def get_stats(self) -> Dict:
        total_size = sum(info.size for info in self.file_infos.values())
        total_chunks = sum(info.chunks for info in self.file_infos.values())
        return {
            'files': len(self.file_infos),
            'chunks': total_chunks,
            'size': total_size,
            'size_formatted': FileUtils.format_size(total_size),
            'folder': self.current_folder
        }
    
    def clear_cache(self) -> TaskResult:
        if os.path.exists(self.cache_path):
            shutil.rmtree(self.cache_path, ignore_errors=True)
        self._search_cache.clear()
        return TaskResult(True, "ìºì‹œ ì‚­ì œ ì™„ë£Œ")
    
    def cleanup(self):
        self.documents.clear()
        self.doc_meta.clear()
        if self.bm25:
            self.bm25.clear()
        self._search_cache.clear()
        try:
            self._executor.shutdown(wait=False)
        except Exception as e:
            logger.debug(f"Executor shutdown error: {e}")
        gc.collect()


# ============================================================================
# Flask ì• í”Œë¦¬ì¼€ì´ì…˜
# ============================================================================
# NumPy í˜¸í™˜ JSON Encoder
class CustomJSONEncoder(json.JSONEncoder):
    def default(self, obj):
        try:
            import numpy as np
            if isinstance(obj, (np.integer, np.floating, np.bool_)):
                return obj.item()
            if isinstance(obj, np.ndarray):
                return obj.tolist()
        except ImportError:
            pass
        return super().default(obj)

app = Flask(__name__)
app.json_encoder = CustomJSONEncoder  # Flask < 2.2 í˜¸í™˜
app.config['MAX_CONTENT_LENGTH'] = AppConfig.MAX_CONTENT_LENGTH
app.config['JSON_AS_ASCII'] = False
app.secret_key = os.urandom(24)
CORS(app, supports_credentials=True)

# Flask > 2.2 í˜¸í™˜ì„ ìœ„í•œ Provider ì„¤ì • (ì„ íƒì )
try:
    from flask.json.provider import DefaultJSONProvider
    class CustomJSONProvider(DefaultJSONProvider):
        def default(self, obj):
            try:
                import numpy as np
                if isinstance(obj, (np.integer, np.floating, np.bool_)):
                    return obj.item()
                if isinstance(obj, np.ndarray):
                    return obj.tolist()
            except ImportError:
                pass
            return super().default(obj)
    app.json = CustomJSONProvider(app)
except ImportError:
    pass  # êµ¬ë²„ì „ FlaskëŠ” json_encoderë§Œ ì‚¬ìš©


# ì „ì—­ QA ì‹œìŠ¤í…œ
qa_system = RegulationQASystem()

# ì—…ë¡œë“œ í´ë” ì„¤ì •
UPLOAD_DIR = os.path.join(get_app_directory(), AppConfig.UPLOAD_FOLDER)
os.makedirs(UPLOAD_DIR, exist_ok=True)


# ============================================================================
# ì—ëŸ¬ í•¸ë“¤ëŸ¬
# ============================================================================
@app.errorhandler(404)
def not_found(e):
    if request.path.startswith('/api/'):
        return jsonify({'success': False, 'message': 'API ì—”ë“œí¬ì¸íŠ¸ë¥¼ ì°¾ì„ ìˆ˜ ì—†ìŠµë‹ˆë‹¤'}), 404
    return render_template('index.html'), 404

@app.errorhandler(500)
def server_error(e):
    logger.error(f"ì„œë²„ ë‚´ë¶€ ì˜¤ë¥˜: {e}")
    if request.path.startswith('/api/'):
        return jsonify({'success': False, 'message': 'ì„œë²„ ë‚´ë¶€ ì˜¤ë¥˜ê°€ ë°œìƒí–ˆìŠµë‹ˆë‹¤', 'error': str(e)}), 500
    return "ì„œë²„ ì˜¤ë¥˜ë°œìƒ", 500

@app.errorhandler(Exception)
def handle_exception(e):
    logger.error(f"ì˜ˆì™¸ ë°œìƒ: {e}")
    import traceback
    logger.error(traceback.format_exc())
    if request.path.startswith('/api/'):
        return jsonify({'success': False, 'message': f'ì˜¤ë¥˜ê°€ ë°œìƒí–ˆìŠµë‹ˆë‹¤: {str(e)}'}), 500
    return str(e), 500



# ============================================================================
# API ë¼ìš°íŠ¸
# ============================================================================
@app.route('/')
def index():
    """ë©”ì¸ ê²€ìƒ‰ í˜ì´ì§€"""
    return render_template('index.html')


@app.route('/admin')
def admin():
    """ê´€ë¦¬ì í˜ì´ì§€"""
    # ì„¸ì…˜ ê¸°ë°˜ ë‹¨ìˆœ ì¸ì¦ ì²´í¬
    # GUI ì„¤ì • ê´€ë¦¬ìì—ì„œ ë¹„ë°€ë²ˆí˜¸ê°€ ì„¤ì •ë˜ì–´ ìˆëŠ”ì§€ í™•ì¸
    settings_manager = get_settings_manager()
    if settings_manager and settings_manager.has_admin_password():
        # ë¹„ë°€ë²ˆí˜¸ê°€ ì„¤ì •ë˜ì–´ ìˆëŠ”ë° ì„¸ì…˜ì— ë¡œê·¸ì¸ì´ ì•ˆë˜ì–´ ìˆë‹¤ë©´
        if not session.get('admin_authenticated'):
            # API ìš”ì²­ì´ë©´ 401, ë¸Œë¼ìš°ì € ì ‘ê·¼ì´ë©´ JSì—ì„œ ì²˜ë¦¬í•˜ë„ë¡ í…œí”Œë¦¿ ë Œë”ë§í•˜ë˜
            # í…œí”Œë¦¿ ë‚´ì—ì„œ ì¸ì¦ ëª¨ë‹¬ì„ ë„ìš°ëŠ” ë°©ì‹ì€ ì´ë¯¸ êµ¬í˜„ë˜ì–´ ìˆìŒ.
            # í•˜ì§€ë§Œ "ì•„ë¬´ëŸ° ì œí•œ ì—†ì´ ì§„ì… ê°€ëŠ¥"í•˜ë‹¤ëŠ” ì‚¬ìš©ì í”¼ë“œë°±ì„ ë°˜ì˜í•˜ì—¬
            # ì—¬ê¸°ì„œëŠ” í…œí”Œë¦¿ì„ ê·¸ëŒ€ë¡œ ì£¼ë˜, app.jsì—ì„œ ì´ˆê¸° ë¡œë“œ ì‹œ ì¸ì¦ì„ ê°•ì œí•˜ë„ë¡ í•¨.
            # ë˜ëŠ” ì—¬ê¸°ì„œ ë°”ë¡œ ì ‘ê·¼ ê±°ë¶€(403)ë¥¼ í•  ìˆ˜ë„ ìˆìŒ.
            # ì‚¬ìš©ì ìš”ì²­: "ì‹¤ì œ ì›¹ì—ì„œëŠ” ì•„ë¬´ëŸ° ì œí•œì—†ì´ ê´€ë¦¬ì ëª¨ë“œë¡œ ì§„ì…ì´ ê°€ëŠ¥"
            # -> JS ì¸ì¦ ì²´í¬ê°€ ìš°íšŒë  ìˆ˜ ìˆìœ¼ë¯€ë¡œ ì„œë²„ì‚¬ì´ë“œ ì²´í¬ê°€ í•„ìš”í•˜ì§€ë§Œ,
            #    í˜„ì¬ êµ¬ì¡°ìƒ ë¡œê·¸ì¸ í˜ì´ì§€ê°€ ë³„ë„ë¡œ ì—†ìœ¼ë¯€ë¡œ, 
            #    JSì—ì„œ ë¹„ë°€ë²ˆí˜¸ ëª¨ë‹¬ì„ ë„ìš°ëŠ” ê²ƒì´ UXìƒ ì¢‹ìŒ.
            #    ë‹¨, API ìš”ì²­ì— ëŒ€í•´ì„œëŠ” ì² ì €íˆ ë§‰ì•„ì•¼ í•¨.
            pass

    return render_template('admin.html')


@app.route('/api/models', methods=['GET'])
def api_get_models():
    """ì‚¬ìš© ê°€ëŠ¥í•œ ëª¨ë¸ ëª©ë¡ ë°˜í™˜"""
    return jsonify({
        'success': True,
        'models': list(AppConfig.AVAILABLE_MODELS.keys()),  # í‚¤ ëª©ë¡ë§Œ ë°˜í™˜
        'current': getattr(qa_system, 'model_name', AppConfig.DEFAULT_MODEL)
    })






# ============================================================================
# ê´€ë¦¬ì ì¸ì¦ API
# ============================================================================
# ì „ì—­ ì„¤ì • ê´€ë¦¬ì ìºì‹œ (ìˆœí™˜ ì°¸ì¡° ë°©ì§€)
_settings_manager_instance = None

def get_settings_manager():
    """GUIì˜ SettingsManager ê°€ì ¸ì˜¤ê¸° (ìˆœí™˜ ì°¸ì¡° ë°©ì§€)"""
    global _settings_manager_instance
    
    if _settings_manager_instance is not None:
        return _settings_manager_instance
    
    try:
        # server_guiê°€ ì´ë¯¸ import ë˜ì—ˆëŠ”ì§€ í™•ì¸
        import sys
        if 'server_gui' in sys.modules:
            _settings_manager_instance = sys.modules['server_gui'].settings_manager
            return _settings_manager_instance
        return None
    except Exception as e:
        logger.debug(f"SettingsManager ì ‘ê·¼ ì‹¤íŒ¨: {e}")
        return None


@app.route('/api/admin/check')
def api_admin_check():
    """ê´€ë¦¬ì ì¸ì¦ ìƒíƒœ í™•ì¸"""
    from flask import session
    
    sm = get_settings_manager()
    if sm is None or not sm.has_admin_password():
        # ë¹„ë°€ë²ˆí˜¸ ë¯¸ì„¤ì • - ì¸ì¦ ë¶ˆí•„ìš”
        return jsonify({'success': True, 'authenticated': True, 'required': False})
    
    # ì„¸ì…˜ í™•ì¸
    is_auth = session.get('admin_authenticated', False)
    return jsonify({'success': True, 'authenticated': is_auth, 'required': True})


@app.route('/api/admin/auth', methods=['POST'])
def api_admin_auth():
    """ê´€ë¦¬ì ë¹„ë°€ë²ˆí˜¸ ì¸ì¦"""
    from flask import session
    
    sm = get_settings_manager()
    if sm is None:
        return jsonify({'success': True, 'message': 'ì¸ì¦ ë¶ˆí•„ìš”'})
    
    data = request.get_json()
    password = data.get('password', '') if data else ''
    
    if sm.verify_admin_password(password):
        session['admin_authenticated'] = True
        return jsonify({'success': True, 'message': 'ì¸ì¦ ì„±ê³µ'})
    else:
        return jsonify({'success': False, 'message': 'ë¹„ë°€ë²ˆí˜¸ê°€ ì¼ì¹˜í•˜ì§€ ì•ŠìŠµë‹ˆë‹¤'}), 401


@app.route('/api/admin/logout', methods=['POST'])
def api_admin_logout():
    """ê´€ë¦¬ì ë¡œê·¸ì•„ì›ƒ"""
    from flask import session
    session.pop('admin_authenticated', None)
    return jsonify({'success': True, 'message': 'ë¡œê·¸ì•„ì›ƒ ì™„ë£Œ'})


@app.route('/api/status')
def api_status():
    """ì„œë²„ ìƒíƒœ ì¡°íšŒ (ì‹œìŠ¤í…œ ë©”íŠ¸ë¦­ í¬í•¨)"""
    # ì‹œìŠ¤í…œ ë©”íŠ¸ë¦­ ìˆ˜ì§‘
    try:
        import psutil
        cpu_percent = psutil.cpu_percent(interval=0.1)
        memory_percent = psutil.virtual_memory().percent
    except ImportError:
        cpu_percent = None
        memory_percent = None
    
    return jsonify({
        'success': True,
        'ready': qa_system.is_ready,
        'loading': qa_system.is_loading,
        'progress': qa_system.load_progress,
        'error': qa_system.load_error,  # ì˜¤ë¥˜ ë©”ì‹œì§€ ì¶”ê°€
        'model': qa_system.model_name,
        'stats': qa_system.get_stats() if qa_system.is_ready else None,
        # ì‹œìŠ¤í…œ ë©”íŠ¸ë¦­
        'cpu_percent': cpu_percent,
        'memory_percent': memory_percent,
        # ê²€ìƒ‰ í ìƒíƒœ
        'search_queue': search_queue.get_stats()
    })


@app.route('/api/health')
def api_health():
    """í—¬ìŠ¤ì²´í¬ ì—”ë“œí¬ì¸íŠ¸"""
    try:
        import psutil
        cpu_percent = psutil.cpu_percent(interval=0.1)
        memory = psutil.virtual_memory()
        memory_percent = memory.percent
    except ImportError:
        cpu_percent = None
        memory_percent = None
    
    return jsonify({
        'status': 'healthy' if qa_system.is_ready else 'initializing',
        'ready': qa_system.is_ready,
        'model_loaded': qa_system.embedding_model is not None,
        'documents_loaded': qa_system.vector_store is not None,
        'document_count': len(qa_system.documents) if qa_system.documents else 0,
        'file_count': len(qa_system.file_infos),
        'cpu_percent': cpu_percent,
        'memory_percent': memory_percent,
        'version': AppConfig.APP_VERSION
    })


@app.route('/api/upload', methods=['POST'])
def api_upload():
    """íŒŒì¼ ì—…ë¡œë“œ ë° ì²˜ë¦¬ (ë™ì‹œ ì—…ë¡œë“œ ë³´í˜¸)"""
    if not qa_system.is_ready:
        return jsonify({'success': False, 'message': 'ì„œë²„ê°€ ì¤€ë¹„ë˜ì§€ ì•Šì•˜ìŠµë‹ˆë‹¤'}), 503
    
    if 'files' not in request.files:
        return jsonify({'success': False, 'message': 'íŒŒì¼ì´ ì—†ìŠµë‹ˆë‹¤'}), 400
    
    files = request.files.getlist('files')
    if not files or files[0].filename == '':
        return jsonify({'success': False, 'message': 'íŒŒì¼ì´ ì„ íƒë˜ì§€ ì•Šì•˜ìŠµë‹ˆë‹¤'}), 400
    
    # íŒŒì¼ ì €ì¥ (ë½ ì‚¬ìš©)
    saved_files = []
    save_errors = []
    
    with file_operation_lock:
        for file in files:
            if file and FileUtils.allowed_file(file.filename):
                try:
                    filename = secure_filename(file.filename)
                    # í•œê¸€ íŒŒì¼ëª… ë³´ì¡´
                    if filename != file.filename:
                        filename = file.filename.replace('/', '_').replace('\\', '_').replace('..', '')
                    
                    # íŒŒì¼ëª… ì¤‘ë³µ ì²˜ë¦¬
                    base_filepath = os.path.join(UPLOAD_DIR, filename)
                    filepath = base_filepath
                    counter = 1
                    while os.path.exists(filepath) and counter < 100:
                        name, ext = os.path.splitext(filename)
                        filepath = os.path.join(UPLOAD_DIR, f"{name}_{counter}{ext}")
                        counter += 1
                    
                    file.save(filepath)
                    saved_files.append(filepath)
                    logger.info(f"íŒŒì¼ ì—…ë¡œë“œ: {os.path.basename(filepath)}")
                except Exception as e:
                    save_errors.append(f"{file.filename}: {e}")
                    logger.error(f"íŒŒì¼ ì €ì¥ ì‹¤íŒ¨: {file.filename} - {e}")
    
    if not saved_files:
        error_msg = 'ì§€ì›ë˜ëŠ” íŒŒì¼ì´ ì—†ìŠµë‹ˆë‹¤'
        if save_errors:
            error_msg += f' (ì˜¤ë¥˜: {", ".join(save_errors[:3])})'
        return jsonify({'success': False, 'message': error_msg}), 400
    
    # ë¬¸ì„œ ì²˜ë¦¬
    def progress_cb(percent, msg):
        logger.info(f"ì²˜ë¦¬ ì§„í–‰: {percent}% - {msg}")
    
    try:
        result = qa_system.process_documents(UPLOAD_DIR, saved_files, progress_cb)
        
        return jsonify({
            'success': result.success,
            'message': result.message,
            'data': result.data,
            'failed': result.failed_items,
            'save_errors': save_errors if save_errors else None
        })
    except Exception as e:
        logger.error(f"ë¬¸ì„œ ì²˜ë¦¬ ì¤‘ ì˜¤ë¥˜: {e}")
        return jsonify({
            'success': False,
            'message': f'ë¬¸ì„œ ì²˜ë¦¬ ì¤‘ ì˜¤ë¥˜: {e}',
            'saved_count': len(saved_files)
        }), 500



@app.route('/api/process', methods=['POST'])
def api_process():
    """ì—…ë¡œë“œëœ ëª¨ë“  íŒŒì¼ ì¬ì²˜ë¦¬"""
    if not qa_system.is_ready:
        return jsonify({'success': False, 'message': 'ì„œë²„ê°€ ì¤€ë¹„ë˜ì§€ ì•Šì•˜ìŠµë‹ˆë‹¤'}), 503
    
    if not os.path.exists(UPLOAD_DIR):
        return jsonify({'success': False, 'message': 'ì—…ë¡œë“œ í´ë”ê°€ ì—†ìŠµë‹ˆë‹¤'}), 400
    
    files = [
        os.path.join(UPLOAD_DIR, f) for f in os.listdir(UPLOAD_DIR)
        if FileUtils.allowed_file(f)
    ]
    
    if not files:
        return jsonify({'success': False, 'message': 'ì²˜ë¦¬í•  íŒŒì¼ì´ ì—†ìŠµë‹ˆë‹¤'}), 400
    
    result = qa_system.process_documents(UPLOAD_DIR, files, None)
    
    return jsonify({
        'success': result.success,
        'message': result.message,
        'data': result.data,
        'failed': result.failed_items
    })


@app.route('/api/search', methods=['POST'])
def api_search():
    """ê²€ìƒ‰ ìˆ˜í–‰"""
    start_time = time.time()
    
    # Rate Limiting ì²´í¬
    client_ip = request.remote_addr or '127.0.0.1'
    if not rate_limiter.is_allowed(client_ip):
        return jsonify({
            'success': False, 
            'message': 'ìš”ì²­ì´ ë„ˆë¬´ ë§ìŠµë‹ˆë‹¤. ì ì‹œ í›„ ë‹¤ì‹œ ì‹œë„í•´ì£¼ì„¸ìš”.',
            'retry_after': 60,
            'remaining_requests': 0
        }), 429
    
    # ê²€ìƒ‰ í ìŠ¬ë¡¯ íšë“
    if not search_queue.acquire(timeout=AppConfig.REQUEST_TIMEOUT):
        return jsonify({
            'success': False,
            'message': 'ì„œë²„ê°€ ë°”ì©ë‹ˆë‹¤. ì ì‹œ í›„ ë‹¤ì‹œ ì‹œë„í•´ì£¼ì„¸ìš”.',
            'queue_stats': search_queue.get_stats()
        }), 503
    
    try:
        if not qa_system.is_ready:
            return jsonify({'success': False, 'message': 'ì„œë²„ê°€ ì¤€ë¹„ë˜ì§€ ì•Šì•˜ìŠµë‹ˆë‹¤'}), 503
        
        if not qa_system.vector_store:
            return jsonify({'success': False, 'message': 'ë¬¸ì„œê°€ ë¡œë“œë˜ì§€ ì•Šì•˜ìŠµë‹ˆë‹¤'}), 400
        
        data = request.get_json()
        if not data or 'query' not in data:
            return jsonify({'success': False, 'message': 'ê²€ìƒ‰ì–´ê°€ í•„ìš”í•©ë‹ˆë‹¤'}), 400
        
        query = data.get('query', '').strip()
        k = min(data.get('k', AppConfig.DEFAULT_SEARCH_RESULTS), AppConfig.MAX_SEARCH_RESULTS)
        hybrid = data.get('hybrid', True)
        highlight = data.get('highlight', True)
        filter_file = data.get('filter_file', None)  # íŒŒì¼ í•„í„°
        sort_by = data.get('sort_by', 'relevance')  # ì •ë ¬ ë°©ì‹
        
        result = qa_system.search(query, k, hybrid, filter_file, sort_by)
        
        # ê²€ìƒ‰ ì„±ê³µ ì‹œ íˆìŠ¤í† ë¦¬ì— ì¶”ê°€
        if result.success and query:
            qa_system._search_history.add(query)
        
        # í•˜ì´ë¼ì´íŒ… ì ìš©
        results_data = result.data if result.success else []
        if highlight and results_data:
            for item in results_data:
                item['content_highlighted'] = TextHighlighter.highlight(item['content'], query)
        
        # ì‘ë‹µ ì‹œê°„ ê³„ì‚°
        response_time_ms = round((time.time() - start_time) * 1000, 2)
        
        return jsonify({
            'success': result.success,
            'message': result.message,
            'results': results_data,
            'query': query,
            'response_time_ms': response_time_ms,
            'result_count': len(results_data)
        })
    except Exception as e:
        logger.error(f"ê²€ìƒ‰ ì¤‘ ì˜¤ë¥˜ ë°œìƒ: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({'success': False, 'message': f'ê²€ìƒ‰ ì²˜ë¦¬ ì¤‘ ì˜¤ë¥˜ê°€ ë°œìƒí–ˆìŠµë‹ˆë‹¤: {e}'}), 500
    finally:
        # ê²€ìƒ‰ í ìŠ¬ë¡¯ ë°˜í™˜
        search_queue.release()



@app.route('/api/search/history')
def api_search_history():
    """ê²€ìƒ‰ íˆìŠ¤í† ë¦¬ ì¡°íšŒ"""
    limit = request.args.get('limit', 10, type=int)
    recent = qa_system._search_history.get_recent(limit)
    popular = qa_system._search_history.get_popular(limit)
    
    return jsonify({
        'success': True,
        'recent': recent,
        'popular': [{'query': q, 'count': c} for q, c in popular]
    })


@app.route('/api/search/suggest')
def api_search_suggest():
    """ê²€ìƒ‰ì–´ ìë™ì™„ì„± ì œì•ˆ"""
    prefix = request.args.get('q', '').strip()
    limit = request.args.get('limit', 8, type=int)
    
    if len(prefix) < 1:
        return jsonify({'success': True, 'suggestions': []})
    
    suggestions = []
    
    # 1. ê²€ìƒ‰ íˆìŠ¤í† ë¦¬ì—ì„œ ë§¤ì¹­
    history_suggestions = qa_system._search_history.suggest(prefix, limit)
    suggestions.extend(history_suggestions)
    
    # 2. ë¬¸ì„œ í‚¤ì›Œë“œì—ì„œ ë§¤ì¹­
    if len(suggestions) < limit:
        keywords = qa_system.get_keywords()
        prefix_lower = prefix.lower()
        for kw in keywords:
            if kw.lower().startswith(prefix_lower) and kw not in suggestions:
                suggestions.append(kw)
                if len(suggestions) >= limit:
                    break
    
    return jsonify({
        'success': True,
        'suggestions': suggestions[:limit]
    })


@app.route('/api/stats/search')
def api_search_stats():
    """ê²€ìƒ‰ í†µê³„ ì¡°íšŒ"""
    limit = request.args.get('limit', 10, type=int)
    
    # ì¸ê¸° ê²€ìƒ‰ì–´
    popular = qa_system._search_history.get_popular(limit)
    
    # ìµœê·¼ ê²€ìƒ‰ì–´
    recent = qa_system._search_history.get_recent(limit)
    
    # ì´ ê²€ìƒ‰ íšŸìˆ˜
    total_searches = sum(count for _, count in popular) if popular else 0
    
    # íŒŒì¼ë³„ ê²€ìƒ‰ ë¹ˆë„ (ê²€ìƒ‰ ê²°ê³¼ì—ì„œ ìì£¼ ë‚˜ì˜¤ëŠ” íŒŒì¼)
    file_stats = []
    for info in qa_system.file_infos.values():
        file_stats.append({
            'name': info.name,
            'chunks': info.chunks,
            'status': info.status.value
        })
    
    return jsonify({
        'success': True,
        'total_searches': total_searches,
        'popular_queries': [{'query': q, 'count': c} for q, c in popular],
        'recent_queries': recent,
        'file_stats': sorted(file_stats, key=lambda x: x['chunks'], reverse=True)[:limit]
    })


@app.route('/api/files/names')
def api_file_names():
    """íŒŒì¼ëª… ëª©ë¡ë§Œ ë°˜í™˜ (ê²€ìƒ‰ í•„í„°ìš©)"""
    names = [info.name for info in qa_system.file_infos.values()]
    return jsonify({
        'success': True,
        'files': sorted(names)
    })


@app.route('/api/files')
def api_files():
    """ë¡œë“œëœ íŒŒì¼ ëª©ë¡"""
    return jsonify({
        'success': True,
        'files': qa_system.get_file_infos(),
        'stats': qa_system.get_stats()
    })


@app.route('/api/cache', methods=['DELETE'])
def api_clear_cache():
    """ìºì‹œ ì‚­ì œ"""
    result = qa_system.clear_cache()
    return jsonify({'success': result.success, 'message': result.message})


@app.route('/api/models')
def api_models():
    """ì‚¬ìš© ê°€ëŠ¥í•œ ëª¨ë¸ ëª©ë¡"""
    return jsonify({
        'success': True,
        'models': list(AppConfig.AVAILABLE_MODELS.keys()),
        'current': qa_system.model_name
    })


@app.route('/api/files/<filename>', methods=['DELETE'])
def api_delete_file(filename):
    """ê°œë³„ íŒŒì¼ ì‚­ì œ (ìŠ¤ë ˆë“œ ì•ˆì „)"""
    if not qa_system.is_ready:
        return jsonify({'success': False, 'message': 'ì„œë²„ê°€ ì¤€ë¹„ë˜ì§€ ì•Šì•˜ìŠµë‹ˆë‹¤'}), 503
    
    # ê²½ë¡œ ê²€ì¦ (ê²½ë¡œ íƒìƒ‰ ê³µê²© ë°©ì§€)
    safe_filename = secure_filename(filename)
    if safe_filename != filename and not filename.replace(' ', '_') == safe_filename:
        # í•œê¸€ íŒŒì¼ëª… ì²˜ë¦¬
        safe_filename = filename.replace('/', '_').replace('\\', '_').replace('..', '')
    
    filepath = os.path.join(UPLOAD_DIR, safe_filename)
    
    # íŒŒì¼ ì¡´ì¬ í™•ì¸
    if not os.path.exists(filepath):
        return jsonify({'success': False, 'message': 'íŒŒì¼ì„ ì°¾ì„ ìˆ˜ ì—†ìŠµë‹ˆë‹¤'}), 404
    
    # íŒŒì¼ ì‚­ì œ (ë½ ì‚¬ìš©)
    with file_operation_lock:
        try:
            # íŒŒì¼ ì‚­ì œ
            os.remove(filepath)
            logger.info(f"íŒŒì¼ ì‚­ì œ: {safe_filename}")
            
            # ì¸ë±ìŠ¤ì—ì„œ í•´ë‹¹ íŒŒì¼ ê´€ë ¨ ë°ì´í„° ì œê±°
            if filepath in qa_system.file_infos:
                del qa_system.file_infos[filepath]
            
            # ê²€ìƒ‰ ìºì‹œ ë¬´íš¨í™”
            qa_system._search_cache.clear()
            qa_system._keyword_cache.clear()
            
            # ë‚¨ì€ íŒŒì¼ë¡œ ì¸ë±ìŠ¤ ì¬êµ¬ì„± í•„ìš” ì•Œë¦¼
            remaining_files = [
                os.path.join(UPLOAD_DIR, f) for f in os.listdir(UPLOAD_DIR)
                if FileUtils.allowed_file(f)
            ]
            
            return jsonify({
                'success': True,
                'message': f'{safe_filename} ì‚­ì œ ì™„ë£Œ',
                'remaining_files': len(remaining_files),
                'reindex_required': True  # í”„ë¡ íŠ¸ì—”ë“œì—ì„œ ì¬ì²˜ë¦¬ ì•ˆë‚´
            })
            
        except OSError as e:
            logger.error(f"íŒŒì¼ ì‚­ì œ ì‹¤íŒ¨: {e}")
            return jsonify({'success': False, 'message': f'ì‚­ì œ ì‹¤íŒ¨: {e}'}), 500
        except Exception as e:
            logger.error(f"íŒŒì¼ ì‚­ì œ ì¤‘ ì˜ˆì™¸: {e}")
            return jsonify({'success': False, 'message': f'ì˜ˆì™¸ ë°œìƒ: {e}'}), 500



@app.route('/api/files/<filename>/preview')
def api_file_preview(filename):
    """íŒŒì¼ ë‚´ìš© ë¯¸ë¦¬ë³´ê¸°"""
    if not qa_system.is_ready:
        return jsonify({'success': False, 'message': 'ì„œë²„ê°€ ì¤€ë¹„ë˜ì§€ ì•Šì•˜ìŠµë‹ˆë‹¤'}), 503
    
    # ê²½ë¡œ ê²€ì¦
    safe_filename = filename.replace('/', '_').replace('\\', '_').replace('..', '')
    filepath = os.path.join(UPLOAD_DIR, safe_filename)
    
    if not os.path.exists(filepath):
        return jsonify({'success': False, 'message': 'íŒŒì¼ì„ ì°¾ì„ ìˆ˜ ì—†ìŠµë‹ˆë‹¤'}), 404
    
    try:
        # íŒŒì¼ ë‚´ìš© ì¶”ì¶œ
        content, error = qa_system.extractor.extract(filepath)
        
        if error:
            return jsonify({
                'success': False,
                'message': f'íŒŒì¼ ì½ê¸° ì‹¤íŒ¨: {error}'
            }), 400
        
        # ë¯¸ë¦¬ë³´ê¸° ê¸¸ì´ ì œí•œ (ê¸°ë³¸ 2000ì)
        max_length = request.args.get('length', 2000, type=int)
        max_length = min(max_length, 5000)  # ìµœëŒ€ 5000ì
        
        preview_content = content[:max_length]
        is_truncated = len(content) > max_length
        
        # íŒŒì¼ ì •ë³´
        file_info = qa_system.file_infos.get(filepath)
        
        return jsonify({
            'success': True,
            'filename': safe_filename,
            'content': preview_content,
            'total_length': len(content),
            'is_truncated': is_truncated,
            'chunks': file_info.chunks if file_info else 0,
            'status': file_info.status.value if file_info else 'unknown'
        })
        
    except Exception as e:
        logger.error(f"ë¯¸ë¦¬ë³´ê¸° ì‹¤íŒ¨: {e}")
        return jsonify({'success': False, 'message': f'ë¯¸ë¦¬ë³´ê¸° ì‹¤íŒ¨: {e}'}), 500


@app.route('/api/files/<filename>/download')
def api_file_download(filename):
    """íŒŒì¼ ë‹¤ìš´ë¡œë“œ"""
    # ê²½ë¡œ ê²€ì¦
    safe_filename = filename.replace('/', '_').replace('\\', '_').replace('..', '')
    filepath = os.path.join(UPLOAD_DIR, safe_filename)
    
    if not os.path.exists(filepath):
        return jsonify({'success': False, 'message': 'íŒŒì¼ì„ ì°¾ì„ ìˆ˜ ì—†ìŠµë‹ˆë‹¤'}), 404
    
    try:
        return send_from_directory(
            UPLOAD_DIR, 
            safe_filename, 
            as_attachment=True,
            download_name=safe_filename
        )
    except Exception as e:
        logger.error(f"íŒŒì¼ ë‹¤ìš´ë¡œë“œ ì‹¤íŒ¨: {e}")
        return jsonify({'success': False, 'message': f'ë‹¤ìš´ë¡œë“œ ì‹¤íŒ¨: {e}'}), 500


@app.route('/api/models', methods=['POST'])
def api_set_model():
    """ëª¨ë¸ ë³€ê²½"""
    # ê´€ë¦¬ì ì¸ì¦ ì²´í¬
    if not session.get('admin_authenticated'):
        settings_manager = get_settings_manager()
        # settings_managerê°€ ìˆê³  íŒ¨ìŠ¤ì›Œë“œê°€ ì„¤ì •ë˜ì–´ ìˆìœ¼ë©´ ì¸ì¦ í•„ìš”
        if settings_manager and settings_manager.has_admin_password():
             return jsonify({'success': False, 'message': 'ê´€ë¦¬ì ê¶Œí•œì´ í•„ìš”í•©ë‹ˆë‹¤'}), 401
             
    data = request.get_json() or {}
    model_name = data.get('model')
    
    if not model_name or model_name not in AppConfig.AVAILABLE_MODELS:
        return jsonify({'success': False, 'message': 'ì˜ëª»ëœ ëª¨ë¸ì…ë‹ˆë‹¤'}), 400
        
    try:
        # ëª¨ë¸ ë³€ê²½ ì‘ì—…
        # 1. ëª¨ë¸ ë¡œë“œ ì‹œë„
        result = qa_system.load_model(model_name)
        
        if result.success:
            # 2. ëª¨ë¸ ë³€ê²½ ì„±ê³µ ì‹œ ìºì‹œ ì´ˆê¸°í™” ë“± ë¶€ê°€ ì‘ì—…? (load_model ë‚´ë¶€ì—ì„œ ì²˜ë¦¬ë¨)
            pass
            
        return jsonify({
            'success': result.success,
            'message': result.message
        })
    except Exception as e:
        logger.error(f"ëª¨ë¸ ë³€ê²½ ì¤‘ ì˜¤ë¥˜: {e}")
        return jsonify({'success': False, 'message': f'ì˜¤ë¥˜ ë°œìƒ: {str(e)}'}), 500


# ============================================================================
# ë©”ì¸ ì‹¤í–‰
# ============================================================================
def load_settings_to_config():
    """ì„¤ì • íŒŒì¼ì—ì„œ AppConfigë¡œ ê°’ ë¡œë“œ"""
    settings_file = os.path.join(get_app_directory(), 'config', 'settings.json')
    if os.path.exists(settings_file):
        try:
            with open(settings_file, 'r', encoding='utf-8') as f:
                settings = json.load(f)
                AppConfig.OFFLINE_MODE = settings.get('offline_mode', False)
                AppConfig.LOCAL_MODEL_PATH = settings.get('local_model_path', '')
                logger.info(f"ğŸ“‹ ì„¤ì • ë¡œë“œ ì™„ë£Œ - ì˜¤í”„ë¼ì¸ ëª¨ë“œ: {AppConfig.OFFLINE_MODE}")
        except Exception as e:
            logger.warning(f"ì„¤ì • íŒŒì¼ ë¡œë“œ ì‹¤íŒ¨ (ê¸°ë³¸ê°’ ì‚¬ìš©): {e}")


def initialize_server():
    """ì„œë²„ ì´ˆê¸°í™” - ëª¨ë¸ ë¡œë“œ"""
    logger.info("=" * 60)
    logger.info(f"ğŸš€ {AppConfig.APP_NAME} {AppConfig.APP_VERSION}")
    logger.info("=" * 60)
    
    # ì„¤ì • íŒŒì¼ì—ì„œ ì˜¤í”„ë¼ì¸ ëª¨ë“œ ì„¤ì • ë¡œë“œ
    load_settings_to_config()
    
    # ì˜¤í”„ë¼ì¸ ëª¨ë“œ ìƒíƒœ ë¡œê¹…
    if AppConfig.OFFLINE_MODE:
        logger.info("ğŸ”’ ì˜¤í”„ë¼ì¸ ëª¨ë“œ í™œì„±í™” - ë¡œì»¬ ëª¨ë¸ë§Œ ì‚¬ìš©")
        if AppConfig.LOCAL_MODEL_PATH:
            logger.info(f"ğŸ“‚ ë¡œì»¬ ëª¨ë¸ ê²½ë¡œ: {AppConfig.LOCAL_MODEL_PATH}")
    else:
        logger.info("ğŸŒ ì˜¨ë¼ì¸ ëª¨ë“œ - í•„ìš” ì‹œ ëª¨ë¸ ë‹¤ìš´ë¡œë“œ")
    
    # ëª¨ë¸ ë¡œë“œ (ì˜¤í”„ë¼ì¸ ëª¨ë“œ ì„¤ì • ì „ë‹¬)
    result = qa_system.load_model(
        AppConfig.DEFAULT_MODEL,
        offline_mode=AppConfig.OFFLINE_MODE,
        local_model_path=AppConfig.LOCAL_MODEL_PATH
    )
    if result.success:
        logger.info(f"âœ… {result.message}")
        
        # ê¸°ì¡´ ì—…ë¡œë“œëœ íŒŒì¼ ìë™ ì²˜ë¦¬
        if os.path.exists(UPLOAD_DIR):
            files = [
                os.path.join(UPLOAD_DIR, f) for f in os.listdir(UPLOAD_DIR)
                if FileUtils.allowed_file(f)
            ]
            if files:
                logger.info(f"ğŸ“‚ ê¸°ì¡´ íŒŒì¼ {len(files)}ê°œ ì²˜ë¦¬ ì¤‘...")
                result = qa_system.process_documents(UPLOAD_DIR, files, None)
                if result.success:
                    logger.info(f"âœ… {result.message}")
                else:
                    logger.warning(f"âš ï¸ {result.message}")
    else:
        logger.error(f"âŒ ëª¨ë¸ ë¡œë“œ ì‹¤íŒ¨: {result.message}")


def graceful_shutdown(signum=None, frame=None):
    """ì„œë²„ ì •ìƒ ì¢…ë£Œ ì²˜ë¦¬"""
    logger.info("ğŸ›‘ ì„œë²„ ì¢…ë£Œ ì¤‘...")
    qa_system.cleanup()
    logger.info("âœ… ì •ìƒ ì¢…ë£Œ ì™„ë£Œ")
    # sys.exit()ëŠ” __main__ì—ì„œë§Œ í˜¸ì¶œ (ê·¸ ì™¸ì—ì„œëŠ” í˜¸ì¶œí•˜ë©´ ì•ˆë¨!)


if __name__ == '__main__':
    # atexit ë“±ë¡ì€ ì§ì ‘ ì‹¤í–‰ ì‹œì—ë§Œ
    atexit.register(graceful_shutdown)
    
    # SIGINT, SIGTERM í•¸ë“¤ëŸ¬ ë“±ë¡
    signal.signal(signal.SIGINT, graceful_shutdown)
    signal.signal(signal.SIGTERM, graceful_shutdown)
    
    # ì„œë²„ ì´ˆê¸°í™” (ë³„ë„ ìŠ¤ë ˆë“œ)
    init_thread = threading.Thread(target=initialize_server, daemon=True)
    init_thread.start()
    
    logger.info(f"ğŸŒ ì„œë²„ ì‹œì‘: http://localhost:{AppConfig.SERVER_PORT}")
    logger.info(f"ğŸ“š ê´€ë¦¬ì í˜ì´ì§€: http://localhost:{AppConfig.SERVER_PORT}/admin")
    logger.info("=" * 60)
    
    # í”„ë¡œë•ì…˜ ì„œë²„ (waitress ê¶Œì¥, ì—†ìœ¼ë©´ Flask ê¸°ë³¸)
    try:
        from waitress import serve
        logger.info("ğŸš€ Waitress í”„ë¡œë•ì…˜ ì„œë²„ë¡œ ì‹¤í–‰")
        serve(app, host=AppConfig.SERVER_HOST, port=AppConfig.SERVER_PORT, threads=8)
    except ImportError:
        logger.warning("âš ï¸ waitress ì—†ìŒ - Flask ê°œë°œ ì„œë²„ ì‚¬ìš© (í”„ë¡œë•ì…˜ì—ì„œëŠ” waitress ê¶Œì¥)")
        app.run(
            host=AppConfig.SERVER_HOST,
            port=AppConfig.SERVER_PORT,
            debug=False,
            threaded=True
        )
